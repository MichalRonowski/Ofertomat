from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import List, Dict
import os

class DOCXGenerator:
    """Klasa do generowania raportów DOCX z ofert"""
    
    def __init__(self):
        self.primary_color = RGBColor(44, 90, 160)  # #2c5aa0
        self.header_color = RGBColor(26, 84, 144)   # #1a5490
    
    def calculate_price(self, purchase_price: float, margin: float, vat_rate: float, quantity: float = 1):
        """
        Kalkuluje ceny
        
        Returns:
            dict z kluczami: net_unit, gross_unit, net_total, vat_amount, gross_total
        """
        # Cena jednostkowa netto sprzedaży
        net_unit = purchase_price * (1 + margin / 100)
        
        # Cena jednostkowa brutto
        gross_unit = net_unit * (1 + vat_rate / 100)
        
        # Wartości dla ilości
        net_total = net_unit * quantity
        vat_amount = net_total * (vat_rate / 100)
        gross_total = net_total + vat_amount
        
        return {
            'net_unit': round(net_unit, 2),
            'gross_unit': round(gross_unit, 2),
            'net_total': round(net_total, 2),
            'vat_amount': round(vat_amount, 2),
            'gross_total': round(gross_total, 2)
        }
    
    def set_cell_background(self, cell, color):
        """Ustawia kolor tła komórki tabeli"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), f"{color.rgb:06x}")
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def generate_offer_docx(self, offer_data: Dict, output_path: str) -> bool:
        """
        Generuje DOCX z ofertą
        
        Args:
            offer_data: Słownik z danymi oferty:
                - title: str (tytuł oferty)
                - date: str (data, opcjonalnie)
                - items: List[Dict] - produkty z polami:
                    - name: str
                    - quantity: float
                    - purchase_price_net: float
                    - vat_rate: float
                    - margin: float (z kategorii)
                    - category_name: str
            output_path: Ścieżka do pliku wyjściowego DOCX
        
        Returns:
            bool - True jeśli sukces
        """
        try:
            # Stwórz katalog jeśli nie istnieje
            os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
            
            # Utwórz dokument
            doc = Document()
            
            # Ustaw marginesy
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
            
            # Tytuł
            title = offer_data.get('title', 'Oferta handlowa')
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = self.header_color
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            date_str = offer_data.get('date', datetime.now().strftime('%d.%m.%Y'))
            date_para = doc.add_paragraph(f"Data: {date_str}")
            date_para.runs[0].font.size = Pt(11)
            
            doc.add_paragraph()  # Spacer
            
            # Pogrupuj produkty po kategoriach
            items_by_category = {}
            for item in offer_data.get('items', []):
                category = item.get('category_name', 'Bez kategorii')
                if category not in items_by_category:
                    items_by_category[category] = []
                items_by_category[category].append(item)
            
            # Dla każdej kategorii
            for category_name, items in sorted(items_by_category.items()):
                # Nagłówek kategorii
                category_para = doc.add_paragraph(category_name)
                category_run = category_para.runs[0]
                category_run.font.size = Pt(14)
                category_run.font.bold = True
                category_run.font.color.rgb = self.primary_color
                
                # Tabela produktów (5 kolumn: Nazwa, Cena netto, J.M., VAT, Cena brutto)
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                # Nagłówki kolumn
                header_cells = table.rows[0].cells
                headers = ['Nazwa', 'Cena netto', 'J.M.', 'VAT', 'Cena brutto']
                
                for i, header_text in enumerate(headers):
                    cell = header_cells[i]
                    cell.text = header_text
                    # Stylizacja nagłówka
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    self.set_cell_background(cell, self.primary_color)
                
                # Wiersze z danymi
                for item in items:
                    prices = self.calculate_price(
                        item['purchase_price_net'],
                        item['margin'],
                        item['vat_rate'],
                        item['quantity']
                    )
                    
                    row_cells = table.add_row().cells
                    
                    # Nazwa
                    row_cells[0].text = item['name']
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                    
                    # Cena netto
                    row_cells[1].text = f"{prices['net_unit']:.2f}"
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(8)
                    
                    # J.M.
                    row_cells[2].text = f"zł/{item.get('unit', 'szt.')}"
                    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_cells[2].paragraphs[0].runs[0].font.size = Pt(8)
                    
                    # VAT
                    row_cells[3].text = f"{item['vat_rate']:.0f}%"
                    row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_cells[3].paragraphs[0].runs[0].font.size = Pt(8)
                    
                    # Cena brutto
                    row_cells[4].text = f"{prices['gross_unit']:.2f} zł"
                    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_cells[4].paragraphs[0].runs[0].font.size = Pt(8)
                
                # Ustaw szerokości kolumn
                table.columns[0].width = Cm(9)
                table.columns[1].width = Cm(2.5)
                table.columns[2].width = Cm(2)
                table.columns[3].width = Cm(1.5)
                table.columns[4].width = Cm(2.5)
                
                doc.add_paragraph()  # Spacer między kategoriami
            
            # Informacja o ważności oferty
            doc.add_paragraph()
            validity_para = doc.add_paragraph()
            validity_run = validity_para.add_run('Oferta ważna w dniu przedstawienia do momentu zmiany cen rynkowych.')
            validity_run.font.size = Pt(8)
            validity_run.font.italic = True
            validity_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Zapisz dokument
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Błąd generowania DOCX: {e}")
            return False
